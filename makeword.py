# # from docx import Document

# # # Create the document
# # doc = Document()

# # # Title
# # doc.add_heading('算法改进详细解析', 0)

# # # Section 1: 幕的设计
# # doc.add_heading('1. 幕的设计', level=1)
# # doc.add_paragraph(
# #     '四幕结构根据曹禺戏剧的典型布局进行划分。算法将每一幕的目标进行结构化：\n'
# #     '- 第 M₁ 幕：潜在矛盾的揭示。情节处于铺垫期，冲突尚未完全爆发。\n'
# #     '- 第 M₂ 幕：冲突的激化。角色之间的对抗加剧，逐步推进至情感高点。\n'
# #     '- 第 M₃ 幕：高潮。所有角色的冲突达到顶点，情感完全爆发。\n'
# #     '- 第 M₄ 幕：解决。角色的命运逐步揭示，冲突结束，剧情收尾。\n'
# #     '通过定义每一幕的核心目标 O(M)，我们可以公式化每幕的设计逻辑：\n'
# #     'O(M₁) = P(z)\n'
# #     'O(M₂) = C(z)\n'
# #     'O(M₃) = E(z)\n'
# #     'O(M₄) = R(z)\n'
# #     '其中，P(z) 为揭示潜在矛盾，C(z) 为冲突激化，E(z) 为情感爆发，R(z) 为冲突解决。'
# # )

# # # Section 2: 情感积累与爆发
# # doc.add_heading('2. 情感积累与爆发', level=1)
# # doc.add_paragraph(
# #     '情感状态的变化通过角色的行为和情境的推进被动态调整。角色的情感状态 Eᵢ(t) 由事件强度 I 逐渐积累，'
# #     '直到达到某一阈值 Tᵢ，触发情感爆发：\n'
# #     'Eᵢ(t+1) = Eᵢ(t) + I\n'
# #     '当 Eᵢ(t) ≥ Tᵢ 时，触发情感爆发，公式为：\n'
# #     'Aᵢ = δ(Eᵢ(t) - Tᵢ)\n'
# #     '其中，Aᵢ 是情感爆发导致的行动，δ(x) 是触发函数。'
# # )

# # # Section 3: 主副线交织
# # doc.add_heading('3. 主副线交织', level=1)
# # doc.add_paragraph(
# #     '在话剧情境设计中，主线 Lₘₐᵢₙ 与副线 Lₛᵤᵦ 的合理交织有助于增强情节复杂性与多样性。'
# #     '算法通过在特定情境中控制主副线角色的交互，来保证各自情节的独立性和少量交叉点。公式化如下：\n'
# #     'Lᵢₙₜ(M₃) = f(Lₘₐᵢₙ, Lₛᵤᵦ)\n'
# #     '其中，Lᵢₙₜ 为主副线的交集，M₃ 表示交集通常在第三幕达到顶点，而函数 f() 表示交互角色的选择。'
# # )

# # # Section 4: 秘密揭露
# # doc.add_heading('4. 秘密揭露', level=1)
# # doc.add_paragraph(
# #     '角色的秘密 Sᵢ 是推动剧情的重要手段。秘密揭露根据每个角色的特质和情感积累进行控制。'
# #     '每一幕中，秘密揭露的概率 P(Sᵢ) 受到幕的进程和情感状态影响：\n'
# #     'P(Sᵢ) = g(M, Eᵢ)\n'
# #     '其中，g() 是揭示函数，结合当前幕 M 和情感状态 Eᵢ，决定是否揭露角色的秘密。'
# #     '通常，秘密在第二或第三幕（M₂, M₃）被揭露。'
# # )

# # # Section 5: 环境变化
# # doc.add_heading('5. 环境变化', level=1)
# # doc.add_paragraph(
# #     '环境变化在情境设计中通过函数 Eₑₙᵥ(t) 动态调整，模拟戏剧冲突外的外部压力。'
# #     '环境状态 Eₑₙᵥ(t) 的改变根据当前剧情进程决定，例如暴雨、地震等自然现象影响角色行为：\n'
# #     'Eₑₙᵥ(t+1) = Eₑₙᵥ(t) + ΔEₛₗₑₙₑ\n'
# #     '当场景数 n 达到阈值时，环境状态随机发生改变。'
# # )

# # # Save the document
# # file_path = './algorithm_improvement_final_v2.docx'
# # doc.save(file_path)

# # file_path



# # from docx import Document

# # def generate_academic_description_word(file_path):
# #     doc = Document()

# #     # 添加标题
# #     doc.add_heading('伏笔与线索在生成情境列表中的添加及呼应机制', 0)

# #     # 添加章节标题与内容
# #     doc.add_heading('1. 伏笔生成与选择', level=1)
# #     doc.add_paragraph(
# #         '伏笔的生成基于提示词 Pₓ（伏笔提示），结合当前情境 Sₜ（当前情境）和历史情境 Hₜ（历史情境），'
# #         '通过提示词动态生成。在每一幕 Mᵢ 的情境 Sₜ 中，算法选择最适合的伏笔类型，并生成对应的内容。\n\n'
# #         '伏笔的生成公式表示为：\n'
# #         'Fₜ = f(Pₓ, Sₜ, Hₜ)\n\n'
# #         '其中，Fₜ 为第 t 场景生成的伏笔，f() 为伏笔生成函数，基于提示词 Pₓ、当前情境 Sₜ 和历史情境 Hₜ。'
# #     )

# #     doc.add_heading('2. 线索生成', level=1)
# #     doc.add_paragraph(
# #         '线索的生成与伏笔相似，但其目的是为后续剧情提供线索，贯穿全剧。线索生成函数 h() 通过提示词 Pₕ '
# #         '生成线索，并确保其贯穿多个情境：\n\n'
# #         'Cₜ = h(Pₕ, Sₜ, Hₜ)\n\n'
# #         '线索提示词 Pₕ 根据剧情的核心冲突和当前情境生成：\n\n'
# #         'Pₕ = k(Sₜ, Hₜ)\n\n'
# #         '贯穿全剧的线索在多个情境中反复出现，提示观众其重要性，并帮助推动剧情发展。'
# #     )

# #     doc.add_heading('3. 伏笔呼应机制', level=1)
# #     doc.add_paragraph(
# #         '在伏笔生成后，需要在后续情境中通过提示词进行呼应和回收，确保之前埋下的伏笔得以揭示。呼应机制通过历史情境和当前情境生成提示词 Pᵣ，'
# #         '然后通过回收函数 fᵣ() 生成呼应内容：\n\n'
# #         'Fᵣₜ = fᵣ(Pᵣ, Fₜ, Sₜ, Hₜ)\n\n'
# #         '其中，Fᵣₜ 是伏笔的呼应内容，提示词 Pᵣ 结合伏笔 Fₜ、当前情境 Sₜ 和历史 Hₜ。\n\n'
# #         '伏笔的回收通常发生在情节的高潮或结局部分：\n\n'
# #         'r(Fₜ) = { 1 if t ≥ tᵣ, 0 otherwise }\n\n'
# #         '其中，tᵣ 为伏笔的回收时间点。'
# #     )

# #     doc.add_heading('4. 算法流程', level=1)
# #     doc.add_paragraph(
# #         '整个伏笔和线索的生成、贯穿及呼应过程包括以下步骤：\n'
# #         '- 生成伏笔 Fₜ 和线索 Cₜ，并在后续情境中回收和贯穿。\n'
# #         '- 使用提示词 Pₓ 和 Pₕ 生成伏笔和线索，提示词选择函数 g() 和 k() 结合当前情境和历史生成合适内容。\n'
# #         '- 在后续情境中通过提示词 Pᵣ 呼应之前的伏笔 Fₜ，确保逻辑连贯。'
# #     )

# #     # 保存到指定路径
# #     doc.save(file_path)
# #     print(f"Word document saved to: {file_path}")



# # # 使用示例
# # file_path = './foreshadowing_clues_academic_description.docx'
# # generate_academic_description_word(file_path)


# from docx import Document

# # def generate_academic_description_word_v2(file_path):
# #     doc = Document()

# #     # 添加标题
# #     doc.add_heading('扩展后的情境生成算法详细说明', 0)

# #     # 生成章节与内容
# #     doc.add_heading('1. 生成角色特质', level=1)
# #     doc.add_paragraph(
# #         '角色的特质（如大五人格模型、光明三角和黑暗三角模型）会直接影响其行为选择和决策逻辑。公式化表示为：\n'
# #         'Aᵢ = f(Tᵢ, Gᵢ, Pᵢ)，其中 Aᵢ 代表角色 i 在特定情境中的行动，Tᵢ 是其性格特质，Gᵢ 为目标，Pᵢ 为计划。'
# #     )

# #     doc.add_heading('2. 生成隐藏信息', level=1)
# #     doc.add_paragraph(
# #         '隐藏信息是推动剧情的重要元素，通常与角色背景、冲突有关。公式为：\n'
# #         'Hᵢ = f(Tᵢ, Cᵢ, Eᵢ)，其中 Hᵢ 为角色 i 的隐藏信息，Cᵢ 为角色冲突，Eᵢ 为情感状态。'
# #     )

# #     doc.add_heading('3. 生成情境事件', level=1)
# #     doc.add_paragraph(
# #         '情境事件由角色的目标、计划、特质以及冲突决定。公式化表示为：\n'
# #         'S(t) = g(Gᵢ, Pᵢ, Tᵢ, Cᵢ, Rᵢ)，S(t) 表示第 t 个情境生成的事件。'
# #     )

# #     doc.add_heading('4. 情感驱动的行动选择', level=1)
# #     doc.add_paragraph(
# #         '角色的行动选择受情感波动的影响，公式为：\n'
# #         'Aᵢ = h(Eᵢ, Gᵢ)，其中 Eᵢ 为情感状态，Aᵢ 为角色行动。'
# #     )

# #     doc.add_heading('5. 计划调整与冲突发展', level=1)
# #     doc.add_paragraph(
# #         '角色的计划会在情境中不断调整，公式化表示为：\n'
# #         'Pᵢ(t+1) = f(Pᵢ(t), Cᵢ(t), Eᵢ(t))，即角色 i 的计划会受到冲突 C 和情感 E 的影响。'
# #     )

# #     doc.add_heading('6. 检测高潮点', level=1)
# #     doc.add_paragraph(
# #         '当情感波动达到一定阈值时，标记情节进入高潮点。公式表示为：\n'
# #         'Climax(t) = 1, 如果 I_C(t) ≥ T_C 或 Eᵢ(t) ≥ T_E。'
# #     )

# #     doc.add_heading('7. 关系变化与情境进展', level=1)
# #     doc.add_paragraph(
# #         '角色的关系会随着情境的发展发生变化，公式为：\n'
# #         'Rᵢⱼ(t+1) = Rᵢⱼ(t) + ΔRᵢⱼ(Cᵢ, Aᵢ)。'
# #     )

# #     doc.add_heading('8. 终止条件', level=1)
# #     doc.add_paragraph(
# #         '当所有角色的主要目标达成时，情境生成终止。公式表示为：\n'
# #         'End = 1, 当 Gᵢ(t) 对所有角色 i 满足时。'
# #     )

# #     # 保存文件
# #     doc.save(file_path)
# #     print(f"Word document saved to: {file_path}")

# # # 使用示例
# # file_path = './academic_scenario_generation_v2.docx'
# # generate_academic_description_word_v2(file_path)

# from docx import Document
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement

# def insert_subscript(paragraph, text):
#     """Helper function to insert subscript into a Word document"""
#     run = paragraph.add_run(text)
#     run.font.subscript = True

# def insert_superscript(paragraph, text):
#     """Helper function to insert superscript into a Word document"""
#     run = paragraph.add_run(text)
#     run.font.superscript = True

# def create_academic_word_doc(file_path):
#     # Create a new Word document
#     doc = Document()

#     # Add the title
#     doc.add_heading('锁闭式结构生成算法的描述', 0)

#     # Section 1: Introduction of climax and conclusion detection
#     doc.add_heading('1. 高潮场景的识别', level=1)
#     paragraph = doc.add_paragraph(
#         '为了识别剧情中的高潮场景，算法通过提示词生成函数 P'
#     )
#     insert_subscript(paragraph, 'climax')
#     paragraph.add_run('，结合当前场景描述 S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run('和历史情境 H')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run('来生成提示信息。每个场景的强度 I(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') 由以下几个因素决定：情感波动 E(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(')（剧烈的情感变化）、冲突强度 C(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(')（角色之间的冲突强度）、行动的不可逆性 A(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(')（角色的关键行动对剧情产生了不可逆影响）。')
    
#     paragraph = doc.add_paragraph('高潮场景的选择公式为：\nS')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' = argmax S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(' ∈ S [P')
#     insert_subscript(paragraph, 'climax')
#     paragraph.add_run('(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run('， H')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') + I(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(')]')

#     # Section 2: Scene adjustment
#     doc.add_heading('2. 场景调整为锁闭式结构', level=1)
#     paragraph = doc.add_paragraph(
#         '在锁闭式结构中，高潮场景 S'
#     )
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run('及其之后的场景 S')
#     insert_subscript(paragraph, 'c+1')
#     paragraph.add_run('...S')
#     insert_subscript(paragraph, 'n')
#     paragraph.add_run('被调整至剧本的开头，以形成强烈的情感冲击。其余场景 S')
#     insert_subscript(paragraph, '1')
#     paragraph.add_run('...S')
#     insert_subscript(paragraph, 'c-1')
#     paragraph.add_run('通过回忆、回溯或角色自白的方式逐步展现，确保剧情逻辑和情感张力的递进。')

#     # Section 3: Climax and plot distribution
#     doc.add_heading('3. 高潮与情节分布', level=1)
#     paragraph = doc.add_paragraph(
#         '算法使用提示词 P'
#     )
#     insert_subscript(paragraph, 'distribution')
#     paragraph.add_run('根据历史情境 H')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run('来调整剧情的分布。高潮后的场景 S')
#     insert_subscript(paragraph, 'c+1')
#     paragraph.add_run('...S')
#     insert_subscript(paragraph, 'n')
#     paragraph.add_run('会被放置在前面，而高潮前的场景 S')
#     insert_subscript(paragraph, '1')
#     paragraph.add_run('...S')
#     insert_subscript(paragraph, 'c-1')
#     paragraph.add_run('通过倒叙、插叙等方式呈现。')

#     # Save the document
#     doc.save(file_path)
#     print(f"Word document saved to: {file_path}")

# # Usage



# # 生成 Word 文档


# # 生成 Word 文档
# file_path = "./closure_structure_academic_description.docx"
# create_academic_word_doc(file_path)



# from docx import Document

# def insert_subscript(paragraph, text):
#     """Helper function to insert subscript into a Word document"""
#     run = paragraph.add_run(text)
#     run.font.subscript = True

# def insert_superscript(paragraph, text):
#     """Helper function to insert superscript into a Word document"""
#     run = paragraph.add_run(text)
#     run.font.superscript = True

# def generate_academic_description_word(file_path):
#     # 创建Word文档
#     doc = Document()

#     # 添加标题
#     doc.add_heading('事件生成的复杂性与情感冲突', 0)

#     # 第1部分：事件生成的定义
#     doc.add_heading('1. 事件生成的定义', level=1)
#     paragraph = doc.add_paragraph('事件生成基于提示词 P')
#     insert_subscript(paragraph, 'event')
#     paragraph.add_run('，结合当前场景 S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(' 和历史情境 H')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run('，为主线或副线中的多个角色生成情感强烈的事件。事件的生成包含角色的行为、冲突的升级、不可逆的动作 A')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(' 和情感波动 E')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run('。事件的强度 I(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') 通过以下公式表示：\n')
#     paragraph = doc.add_paragraph('I(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') = E(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') + C(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') + A(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(')\n')
#     paragraph.add_run('其中，E(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') 表示角色之间的情感波动，C(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') 表示冲突强度，A(S')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(') 表示不可逆动作的强度。')

#     # 第2部分：角色互动与冲突生成
#     doc.add_heading('2. 角色之间的互动与冲突生成', level=1)
#     paragraph = doc.add_paragraph('事件生成中，角色的目标、背景和冲突引导角色的行动。主线角色 R')
#     insert_subscript(paragraph, 'main')
#     paragraph.add_run(' 和副线角色 R')
#     insert_subscript(paragraph, 'sub')
#     paragraph.add_run(' 通过冲突提示词 P')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run(' 和情感提示词 P')
#     insert_subscript(paragraph, 'emotion')
#     paragraph.add_run(' 生成互动。角色选择规则如下：\n')
#     paragraph = doc.add_paragraph('R')
#     insert_subscript(paragraph, 'event')
#     paragraph.add_run(' = R')
#     insert_subscript(paragraph, 'main')
#     paragraph.add_run(' + R')
#     insert_subscript(paragraph, 'sub')
#     paragraph.add_run('\n其中，R')
#     insert_subscript(paragraph, 'event')
#     paragraph.add_run(' 表示事件中的角色集合，R')
#     insert_subscript(paragraph, 'main')
#     paragraph.add_run(' 为主线角色，R')
#     insert_subscript(paragraph, 'sub')
#     paragraph.add_run(' 为副线角色的随机子集。')

#     # 第3部分：情感冲突与反转因素
#     doc.add_heading('3. 事件的情感冲突与反转因素', level=1)
#     paragraph = doc.add_paragraph('为了增加剧情复杂性，事件中可能包含反转或反常行为，由提示词 P')
#     insert_subscript(paragraph, 'surprise')
#     paragraph.add_run(' 引导。角色的情感冲突通过 P')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run(' 生成，保证行动符合剧情发展。\n')
#     paragraph = doc.add_paragraph('E')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run(' = f(P')
#     insert_subscript(paragraph, 'emotion')
#     paragraph.add_run('， P')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run('， R')
#     insert_subscript(paragraph, 'event')
#     paragraph.add_run('， H')
#     insert_subscript(paragraph, 't')
#     paragraph.add_run(')\n其中，E')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run(' 表示情感冲突，f 为事件生成函数，P')
#     insert_subscript(paragraph, 'emotion')
#     paragraph.add_run(' 和 P')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run(' 为情感和冲突提示词。')

#     # 第4部分：事件生成流程
#     doc.add_heading('4. 事件生成的完整流程', level=1)
#     paragraph = doc.add_paragraph(
#         '事件生成流程包含以下步骤：\n'
#         '- 从主线与副线角色中选择参与事件的角色 R'
#     )
#     insert_subscript(paragraph, 'event')
#     paragraph.add_run('\n- 根据提示词 P')
#     insert_subscript(paragraph, 'surprise')
#     paragraph.add_run(' 和 P')
#     insert_subscript(paragraph, 'conflict')
#     paragraph.add_run(' 生成情感冲突与不可逆动作\n'
#                       '- 更新角色关系与情感状态，推动剧情发展至高潮。')

#     # 保存文档
#     doc.save(file_path)
#     print(f"Word document saved to: {file_path}")

# # 使用示例



# file_path = './event_academic_description.docx'
# generate_academic_description_word(file_path)

# from docx import Document

# # Helper function to insert subscript
# def insert_subscript(paragraph, text):
#     run = paragraph.add_run(text)
#     run.font.subscript = True

# # Function to generate academic description
# def generate_closure_structure_description_word(file_path):
#     # Create a new Word document
#     doc = Document()

#     # Add title
#     doc.add_heading('锁闭式结构算法概述', 0)

#     # Part 1: Algorithm for Climax and Conclusion Generation
#     doc.add_heading('1. 高潮与结局的生成算法', level=1)
#     paragraph = doc.add_paragraph('首先，算法从生成高潮 S')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' 和结局 S')
#     insert_subscript(paragraph, 'e')
#     paragraph.add_run(' 开始，通过事件驱动的场景生成，确保高潮场景 S')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' 在情节发展中合理爆发。高潮场景的生成依赖角色的行为 A')
#     insert_subscript(paragraph, 'i')
#     paragraph.add_run('、冲突 C')
#     insert_subscript(paragraph, 'j')
#     paragraph.add_run(' 以及剧情需求 D，通过以下关系表示：\n')
#     paragraph = doc.add_paragraph('S')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' = f(A')
#     insert_subscript(paragraph, 'i')
#     paragraph.add_run('，C')
#     insert_subscript(paragraph, 'j')
#     paragraph.add_run('，D)\n其中，f 是结合所有输入信息的生成函数。')

#     # Part 2: Goal Evaluation for Climax and Conclusion
#     doc.add_heading('2. 目标评估与多次迭代', level=1)
#     paragraph = doc.add_paragraph('高潮生成后，算法使用检查机制 G')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' 来判断高潮是否合格，并通过多次迭代调整直至生成合格的高潮场景。结局 S')
#     insert_subscript(paragraph, 'e')
#     paragraph.add_run(' 的生成使用相似的评估机制 G')
#     insert_subscript(paragraph, 'e')
#     paragraph.add_run('，以确保剧情的闭合性。')

#     # Part 3: Scene Generation in Acts 1 and 2
#     doc.add_heading('3. 前两幕的正序生成', level=1)
#     paragraph = doc.add_paragraph('在生成高潮和结局后，算法正序生成前两幕 S')
#     insert_subscript(paragraph, '1')
#     paragraph.add_run(' 和 S')
#     insert_subscript(paragraph, '2')
#     paragraph.add_run('，逐步展开剧情。每一幕通过阶段性目标检查 G')
#     insert_subscript(paragraph, 'i')
#     paragraph.add_run(' 确认是否达成。如果达成，则进入下一幕生成，否则重新生成当前场景。')

#     # Part 4: Ensuring Connection to Climax
#     doc.add_heading('4. 高潮前的衔接与检查', level=1)
#     paragraph = doc.add_paragraph('为了确保高潮前的情境与高潮 S')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' 自然衔接，算法对高潮前的主线和副线情境进行专门生成，使用提示词 P')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' 进行检查：')
#     paragraph = doc.add_paragraph('G')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(' (S')
#     insert_subscript(paragraph, '2')
#     paragraph.add_run('，S')
#     insert_subscript(paragraph, 'c')
#     paragraph.add_run(') = 1 表示衔接合理，0 表示不合理，需要调整。')

#     # Part 5: Secret Revelation and Exit Mechanism
#     doc.add_heading('5. 秘密揭露与角色退出机制', level=1)
#     paragraph = doc.add_paragraph('算法在适当的场景中揭露角色的秘密 R')
#     insert_subscript(paragraph, 's')
#     paragraph.add_run('，并评估角色是否合理退出。退出由退出机制 E')
#     insert_subscript(paragraph, 'i')
#     paragraph.add_run(' 决定，通过以下函数表示：\n')
#     paragraph = doc.add_paragraph('E')
#     insert_subscript(paragraph, 'i')
#     paragraph.add_run(' (S')
#     insert_subscript(paragraph, 'i')
#     paragraph.add_run(') = 1 表示合理退出，0 表示不退出。')

#     # Save the document
#     doc.save(file_path)
#     print(f"Word document saved to: {file_path}")

# # Example usage
# file_path = './closure_structure_algorithm2.docx'
# generate_closure_structure_description_word(file_path)


# from docx import Document
# from docx.shared import Inches

# # Data for the Word table
# events_data = [
#     {"title": "《原野》", "人物": "仇虎（主角）", "行为": "仇虎为了给父亲复仇，设计杀害焦大星。", "结果": "仇虎成功复仇，但也陷入内心的深刻挣扎，最终未能摆脱仇恨的枷锁。"},
#     {"title": "《原野》", "人物": "金子（悲剧人物）", "行为": "金子爱上仇虎，希望帮助他从复仇的阴影中解脱。", "结果": "金子的爱情未能挽救仇虎，自己也陷入情感上的痛苦。"},
#     {"title": "《原野》", "人物": "焦阎王（反面人物）", "行为": "焦阎王通过暴力和权力压迫村民，掌控他们的命运。", "结果": "他的压迫造成了村民的痛苦，揭示了社会的残酷不公。"},
#     {"title": "《日出》", "人物": "陈白露（主角）", "行为": "陈白露被卷入资本主义腐败与奢靡的生活之中，逐渐失去了自我。", "结果": "她无法从空虚的生活中解脱，最终选择了自我毁灭。"},
#     {"title": "《日出》", "人物": "李石清（悲剧人物）", "行为": "李石清为了改善生活，在潘月亭的公司拼命工作。", "结果": "在无法忍受贫困和压迫下，李石清最终选择自杀，资本主义社会的冷酷无情暴露无遗。"},
#     {"title": "《日出》", "人物": "胡四（配角）", "行为": "作为陈白露的仆人，胡四目睹了上层社会的腐败和堕落。", "结果": "通过观察，胡四认识到自己无法改变命运，阶级压迫感进一步加深。"},
#     {"title": "《北京人》", "人物": "曾文清（主角）", "行为": "曾文清试图维护家族的荣誉，但面对家族的衰败与背叛感到无力。", "结果": "家族内部矛盾加剧，曾文清失去了一切，象征着旧社会家族制度的崩溃。"},
#     {"title": "《北京人》", "人物": "袁任敢（反叛者）", "行为": "袁任敢在科学研究和家庭责任中挣扎，试图打破固有的社会规范。", "结果": "他的反叛引发了家庭和社会的冲突，揭示了知识分子在动荡时代中的精神斗争。"},
#     {"title": "《雷雨》", "人物": "周朴园（反面人物）、鲁侍萍（悲剧人物）", "行为": "周朴园与鲁侍萍重逢，试图用金钱弥补过往的错误。", "结果": "鲁侍萍拒绝补偿，家庭内部的旧怨加剧，揭露了封建社会下的冷酷与压迫。"},
#     {"title": "《雷雨》", "人物": "周萍（主角）、四凤（配角）", "行为": "周萍与四凤发展禁忌恋情，周萍后来得知四凤是他的同父异母的妹妹。", "结果": "两人陷入伦理和情感的困境，最终四凤死亡，周萍自杀，故事以悲剧告终。"}
# ]

# # Create a new Document
# doc = Document()
# doc.add_heading('曹禺剧作事件分析表格', 0)

# # Add table with four columns: Title, 人物, 行为, 结果
# table = doc.add_table(rows=1, cols=4)
# table.autofit = True

# # Define the header row
# header_cells = table.rows[0].cells
# header_cells[0].text = 'Title'
# header_cells[1].text = '人物'
# header_cells[2].text = '行为'
# header_cells[3].text = '结果'

# # Add data rows
# for event in events_data:
#     row_cells = table.add_row().cells
#     row_cells[0].text = event["title"]
#     row_cells[1].text = event["人物"]
#     row_cells[2].text = event["行为"]
#     row_cells[3].text = event["结果"]

# # Save the document
# file_path = "曹禺剧作事件分析表格.docx"
# doc.save(file_path)

# print(f"Document saved as {file_path}")


from docx import Document
from docx.shared import Pt

# Create a new Document
doc = Document()

# Title
doc.add_heading('哲学性角色隐喻构建', 0)

# Table structure
table = doc.add_table(rows=1, cols=5)

# Table header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '案例'
hdr_cells[1].text = '原始对白'
hdr_cells[2].text = '本体提取'
hdr_cells[3].text = '喻体选择与推理'
hdr_cells[4].text = '生成对白'

# Fill in the data for the table
data = [
    ("案例1：哲学性角色隐喻构建", 
     "我们的希望正在逐渐消失。", 
     "核心情感：希望消逝", 
     "情境要素：海上风暴（当前剧情为船只遇险）\n选定喻体：灯塔\n本体-喻体共性：指引性、遥远感、逐渐黯淡\n时间维度：灯塔光束从规律闪烁到间歇性熄灭\n环境变化：风暴潮水位上涨淹没灯塔基座\n角色特质：船长（理性思辨型）", 
     "希望就像风暴中的灯塔——我们计算着它的闪烁频率，却发现潮水早已漫过刻度线。"), 
    ("案例2：情绪化角色隐喻迭代", 
     "我感到如此孤独。", 
     "核心情感：孤独感", 
     "情境要素：深夜实验室（角色独自加班）\n选定喻体：培养皿中的菌落\n本体-喻体共性：封闭环境、增殖中的孤立\n生命状态：菌落从活跃分裂到停止增殖\n空间变化：显微镜视野不断放大却更显逼仄\n角色特质：生物研究员（敏感内倾型）", 
     "我的孤独是第37小时培养皿里的菌落——它们在分裂中彼此远离，直到填满所有空隙却永不相触。"),
    ("案例3：对抗性场景隐喻强化", 
     "甲方：'你的承诺毫无价值！'\n乙方：'我已经尽力了。'", 
     "核心冲突：信任破裂", 
     "情境要素：建筑工地（剧情高潮为桥梁坍塌）\n选定喻体：钢筋应力\n本体-喻体共性：累积损伤、临界断裂\n材料变化：钢筋锈蚀程度随时间递增\n结构响应：细微震颤引发共振效应", 
     "甲方：'你的承诺就像超限应力的钢筋——每个微小裂隙都在说着'我尽力了'，直到整个结构轰然崩塌！'\n乙方：'但你可曾计算过荷载的递增速度？'"),
    ("案例4：命运转折隐喻嵌套", 
     "这就是我的命运。", 
     "核心概念：命运不可逆", 
     "情境要素：化学实验（角色目睹不可逆反应）\n选定喻体：结晶过程\n本体-喻体共性：定向发展、结构固化\n相变过程：过饱和溶液突然析出晶体\n观察视角：显微镜下晶格定向生长", 
     "命运是过饱和溶液里的第一次成核——从那瞬间起，所有分子都朝着既定的晶格方向狂奔，再也回不到混沌的自由。")
]

# Add data rows
for case in data:
    row_cells = table.add_row().cells
    row_cells[0].text = case[0]
    row_cells[1].text = case[1]
    row_cells[2].text = case[2]
    row_cells[3].text = case[3]
    row_cells[4].text = case[4]

# Set table style for three-line format
table.style = 'Table Grid'
# Adjust column width
for col in table.columns:
    for cell in col.cells:
        cell.width = Pt(100)

# Save the document
file_path = "./philosophical_character_metaphor_case_study.docx"
doc.save(file_path)

file_path



二、对比维度与发现
1. 人物一致性：算法对角色动机的捕捉能力
Ibsen剧本：通过高频潜台词（如海妲对乔治的肢体控制）构建角色的权力关系，符黛雅“共同孕育的结晶”等隐喻直指学术政治化议题，依赖深度语义理解实现人物复杂性。

生成剧本：沈寒的玩世不恭（“人生如戏，全靠演技”）与苏青的理想主义（“为自由而战”）通过关键台词反复强化，算法通过情节模板与情绪标签实现角色动机的连贯性。

有效性验证：算法生成的短剧在单一情境下（如咖啡馆对峙）能保持人物行为一致性，但在复杂权力博弈场景中仍需依赖人工调整。

2. 叙事逻辑：结构化与反转设计的适配性
Ibsen剧本：以记者会为叙事支点，通过符黛雅声明、罗艾勒死亡真相揭露等节点实现政治讽刺，非线性冲突依赖因果推理与符号映射。

生成剧本：采用三幕式经典结构（开端-对抗-结局），通过叶辰死亡、林慕白背叛等高频反转推动剧情，算法基于情节弧线（Plot Arc）模型实现节奏控制。

有效性验证：算法在封闭叙事框架内（如短剧的单一主线）展现出高效的情节编排能力，反转密度（平均每场1.2次）符合商业短剧市场需求。

3. 场景转换：空间符号化的技术实现
Ibsen剧本：后台走廊、记者会现场的封闭空间隐喻权力牢笼，依赖场景细节（如手机提示声）激发观众联想。

生成剧本：咖啡馆、江边码头等场景通过“昏黄路灯”“潮湿石板路”等标准化描述符实现氛围渲染，算法基于空间-情绪词典自动匹配环境描写。

有效性验证：算法在场景符号化（如“迷雾”象征混沌）上表现稳定，但对开放式隐喻（如Ibsen剧中“尿尿自由”的社会批判）生成能力有限。